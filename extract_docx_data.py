#!/usr/bin/env python3
"""
Extract student data from .docx files and convert to structured JSON
"""

import json
import re
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

def extract_tables_from_docx(filename):
    """Extract all tables from a .docx file"""
    doc = Document(filename)
    tables_data = []

    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_data.append(row_data)
        tables_data.append(table_data)

    return tables_data

def extract_text_from_docx(filename):
    """Extract all text from a .docx file"""
    doc = Document(filename)
    full_text = []

    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())

    return full_text

def parse_student_record(row):
    """Parse a row of student data into structured format"""
    # Expected columns: Uni, Year, First, Last, Major, Notes
    # Or some variation of this

    if len(row) < 4:
        return None

    # Clean up each field
    cleaned = [cell.strip() for cell in row]

    return {
        'uni': cleaned[0] if len(cleaned) > 0 else '',
        'year': cleaned[1] if len(cleaned) > 1 else '',
        'first': cleaned[2] if len(cleaned) > 2 else '',
        'last': cleaned[3] if len(cleaned) > 3 else '',
        'major': cleaned[4] if len(cleaned) > 4 else '',
        'notes': cleaned[5] if len(cleaned) > 5 else ''
    }

def extract_year_from_filename(filename):
    """Extract year from filename like '2026.docx' """
    match = re.search(r'(\d{4})', filename)
    if match:
        return match.group(1)
    return ''

def process_docx_file(filename):
    """Process a single .docx file and extract student data"""
    print(f"\n=== Processing {filename} ===")

    # Try extracting tables first
    tables = extract_tables_from_docx(filename)

    if tables:
        print(f"Found {len(tables)} table(s)")
        for i, table in enumerate(tables):
            print(f"\nTable {i+1}: {len(table)} rows, {len(table[0]) if table else 0} columns")
            if table and len(table) > 0:
                print(f"Header: {table[0]}")
                if len(table) > 1:
                    print(f"Sample row: {table[1]}")

    # Also extract all text to understand structure
    text = extract_text_from_docx(filename)
    print(f"\nTotal paragraphs: {len(text)}")
    if text:
        print(f"First few lines:")
        for line in text[:10]:
            print(f"  {line}")

    return tables, text

def main():
    """Extract data from all .docx files"""
    docx_files = ['grad.docx', '2026.docx', '2027.docx', '2028.docx', '2029.docx']

    all_students = []

    for filename in docx_files:
        try:
            tables, text = process_docx_file(filename)

            # Extract year from filename as fallback
            file_year = extract_year_from_filename(filename)

            # Process tables
            for table_idx, table in enumerate(tables):
                if not table or len(table) < 2:
                    continue

                # Assume first row is header, skip it
                headers = table[0]
                print(f"\nProcessing table with headers: {headers}")

                for row_idx, row in enumerate(table[1:], start=1):
                    if len(row) < 2 or all(not cell.strip() for cell in row):
                        continue

                    student = parse_student_record(row)
                    if student:
                        # Use file year if no year in data
                        if not student['year'] and file_year:
                            student['year'] = file_year

                        student['source_file'] = filename
                        all_students.append(student)

        except Exception as e:
            print(f"Error processing {filename}: {e}")
            import traceback
            traceback.print_exc()

    print(f"\n=== Summary ===")
    print(f"Total students extracted: {len(all_students)}")

    # Save to JSON
    output_file = 'data/students.json'
    with open(output_file, 'w') as f:
        json.dump(all_students, f, indent=2)

    print(f"Saved to {output_file}")

    # Print sample
    if all_students:
        print(f"\nSample student record:")
        print(json.dumps(all_students[0], indent=2))

    return all_students

if __name__ == '__main__':
    main()
